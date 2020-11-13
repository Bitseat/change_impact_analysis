import docx 
import csv

#doc = docx.Document()
def unicodetoascii(text):

    # if you want to do more preprocessing
    TEXT = (text.
            replace('_x000D_', "").
            replace(' � ', "").
            replace('CSTATYPE', "")

                 )
    return TEXT

with open('all_requirements.csv', newline='') as f:
    csv_reader = csv.reader(f) 

    csv_headers = next(csv_reader)
    csv_cols = len(csv_headers)
    i = 1
    for row in csv_reader:
        doc = docx.Document()
        formatted_output = str(row[1]).replace('\\n', '\n').replace('\\t', '\t')
        formatted_output.encode().decode()
        formatted_output = unicodetoascii(formatted_output)
        formatted_output = formatted_output.replace('b\'', '')
        formatted_output = formatted_output.replace('b\"', '')
        formatted_output = formatted_output.replace('\'', '')
        formatted_output = formatted_output.replace('\"', '')
        formatted_output = formatted_output.replace(':', '')
        formatted_output = formatted_output.replace('\n\n', '\n')

        doc.add_paragraph(formatted_output)
        doc.add_page_break()
        doc.save("docx/"+str(row[0])+".docx")
        i = i+1


